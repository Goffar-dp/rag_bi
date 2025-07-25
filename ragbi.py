import os
import yaml
import hashlib
import fitz  # PyMuPDF
import pytesseract
import re
from datetime import datetime, timedelta
import time
from langdetect import detect, DetectorFactory
from langchain_community.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import tempfile

# Fix for langdetect consistency
DetectorFactory.seed = 0

class GeminiLLM:
    def __init__(self):
        self.last_call = datetime.now() - timedelta(seconds=60)
        self.gemini_model = "gemini-1.5-flash"
        self.rate_limit = 60  # seconds between requests
    
    def generate(self, prompt, is_bengali):
        try:
            # Enforce rate limiting
            now = datetime.now()
            elapsed = (now - self.last_call).total_seconds()
            if elapsed < self.rate_limit:
                time.sleep(self.rate_limit - elapsed)
            
            # Add Bengali instruction if needed
            if is_bengali:
                prompt = f"এই প্রশ্নের উত্তর বাংলায় দিন:\n{prompt}"
            
            gemini = ChatGoogleGenerativeAI(
                model=self.gemini_model,
                temperature=0.2,
                convert_system_message_to_human=True
            )
            
            response = gemini.invoke(prompt)
            self.last_call = datetime.now()
            return response.content
        except Exception as e:
            error_msg = str(e)
            if "quota" in error_msg.lower() or "429" in error_msg:
                if is_bengali:
                    return "দুঃখিত, API কোটা শেষ হয়েছে। অনুগ্রহ করে পরে আবার চেষ্টা করুন।"
                else:
                    return "Sorry, API quota exceeded. Please try again later."
            elif "location is not supported" in error_msg:
                self.gemini_model = "gemini-1.0-pro"
                return self.generate(prompt, is_bengali)
            else:
                if is_bengali:
                    return "দুঃখিত, উত্তর তৈরি করতে সমস্যা হচ্ছে। অনুগ্রহ করে আবার চেষ্টা করুন।"
                else:
                    return "Sorry, I'm having trouble generating a response. Please try again."

class RAGSystem:
    def __init__(self):
        self.vectorstore = None
        self.processed_files = {}
        self.all_docs = []
        self.llm = GeminiLLM()
        self.embedding_model = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )

    def extract_text_with_ocr(self, pdf_path):
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(
                pdf_path, 
                dpi=300, 
                poppler_path="/usr/bin",
                thread_count=4
            )
            
            text = ""
            for image in images:
                page_text = pytesseract.image_to_string(
                    image.convert('L'),
                    lang='ben',
                    config='--psm 6 --oem 3'
                )
                page_text = re.sub(r'[^\u0980-\u09FF\s.,!?;:]+', '', page_text)
                text += f"{page_text}\n\n"
            return text.strip()
        except Exception:
            return ""

    def load_document(self, file_path, file_extension, file_name):
        try:
            if file_extension == ".pdf":
                try:
                    doc = fitz.open(file_path)
                    text = ""
                    for i in range(len(doc)):
                        page = doc.load_page(i)
                        text += page.get_text() + "\n\n"
                    
                    if any('\u0980' <= char <= '\u09FF' for char in text):
                        ocr_text = self.extract_text_with_ocr(file_path)
                        if ocr_text:
                            return [Document(
                                page_content=ocr_text,
                                metadata={"source": file_name, "pages": f"1-{len(doc)}"}
                            )]
                    
                    return [Document(
                        page_content=text,
                        metadata={"source": file_name, "pages": f"1-{len(doc)}"}
                    )]
                except Exception:
                    return []
                
            elif file_extension == ".docx":
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return [Document(
                        page_content=text,
                        metadata={"source": file_name, "pages": "all"}
                    )]
                except Exception:
                    return []
                
            elif file_extension == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                return [Document(
                    page_content=text,
                    metadata={"source": file_name, "pages": "all"}
                )]
            
            else:
                return []
        except Exception:
            return []

    def process_upload(self, file_bytes, file_name):
        file_hash = hashlib.md5(file_bytes).hexdigest()
        if file_hash in self.processed_files:
            return 0  # Already processed
        
        _, file_extension = os.path.splitext(file_name)
        file_extension = file_extension.lower()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        docs = self.load_document(tmp_path, file_extension, file_name)
        os.unlink(tmp_path)
        
        if docs:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", "।", "। ", " ", ""]
            )
            chunks = splitter.split_documents(docs)
            self.all_docs.extend(chunks)
            self.processed_files[file_hash] = file_name
            
            if self.vectorstore is None:
                self.vectorstore = Chroma.from_documents(
                    documents=self.all_docs,
                    embedding=self.embedding_model
                )
            else:
                self.vectorstore.add_documents(chunks)
            
            return len(chunks)
        return 0

    def detect_language(self, text):
        try:
            lang = detect(text)
            return "bn" if lang in ["bn", "as", "or"] else "en"
        except:
            return "en"

    def get_response(self, query):
        if not self.vectorstore:
            return "Please upload documents first", []
        
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 3})
        docs = retriever.get_relevant_documents(query)
        
        lang = self.detect_language(query)
        is_bengali = lang == "bn"
        
        context = "\n\n".join([
            f"Source: {doc.metadata.get('source', 'Unknown')} (Pages: {doc.metadata.get('pages', 'N/A')})\n{doc.page_content}"
            for doc in docs
        ])
        
        if is_bengali:
            prompt = f"প্রসঙ্গ: {context}\nপ্রশ্ন: {query}\nউত্তর: "
        else:
            prompt = f"Context: {context}\nQuestion: {query}\nAnswer:"
        
        answer = self.llm.generate(prompt, is_bengali)
        return answer, docs

    def format_sources(self, docs):
        sources = []
        for doc in docs:
            source = doc.metadata.get("source", "Unknown")
            pages = doc.metadata.get("pages", "N/A")
            content = doc.page_content.strip()
            content = content[:150] + "..." if len(content) > 150 else content
            
            sources.append({
                "Source File": source,
                "Page(s)": pages,
                "Content Excerpt": content
            })
        return sources
